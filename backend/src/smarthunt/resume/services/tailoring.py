from pathlib import Path

from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from smarthunt.ai.service import ai_service
from smarthunt.ai.types import AIRequest
from smarthunt.matching.job_parser import extract_job_skills
from smarthunt.resume.models import TailoredResume
from smarthunt.resume.parser.skills import extract_skills
from smarthunt.resume.storage.storage import STORAGE_DIR

TAILORED_DIR = STORAGE_DIR / "tailored"

# Deliberately asks the model for ONLY a short professional summary, not a
# rewrite of the whole resume — the local model (see settings.ollama_model)
# is small and not trusted to reproduce factual content (dates, companies,
# achievements) from scratch without a real risk of hallucinating details
# on a document that gets submitted to real employers. The real resume text
# is kept 100% verbatim below the AI-written summary.
TAILORED_SUMMARY_PROMPT = """You are a career coach. Based on the real resume and job description \
below, write ONLY a 3-4 sentence professional summary/objective paragraph tailored to this exact \
job, written in the resume owner's voice ("I have..."). Use only real skills/experience that \
actually appear in the resume — do not invent anything. Mention the target role or company by \
name if it appears in the job description. Output only the paragraph, no headers, no markdown, no \
extra commentary.

Resume:
{resume}

Job description:
{job}
"""

# Same reasoning as matching/services/deep_analysis.py and
# cover_letter/service.py: prompt-eval time on the CPU-bound local model
# scales with input length — measured live 2026-08-04, a 3000-char-each
# prompt regularly exceeded the old 150s timeout below. Only the copy fed
# to the AI is capped, the resume actually saved/uploaded below stays whole.
_MAX_CHARS_FOR_AI = 1500


def _truncate_for_ai(text: str) -> str:
    if len(text) <= _MAX_CHARS_FOR_AI:
        return text
    return text[:_MAX_CHARS_FOR_AI] + "\n...(تم اختصار الباقي)"


def _write_docx(path: Path, summary: str, resume_text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(summary)
    doc.add_heading("Resume", level=2)
    for line in resume_text.splitlines():
        doc.add_paragraph(line)
    doc.save(str(path))


async def generate_tailored_resume(
    db: AsyncSession,
    job_id: int,
    resume_text: str,
    job_description: str,
) -> TailoredResume:
    resume_skills = set(extract_skills(resume_text))
    job_skills = set(extract_job_skills(job_description))

    matched = sorted(resume_skills & job_skills)
    missing = sorted(job_skills - resume_skills)
    score = int((len(matched) / len(job_skills)) * 100) if job_skills else 0

    summary = ""
    try:
        ai_response = await ai_service.generate(
            AIRequest(
                prompt=TAILORED_SUMMARY_PROMPT.format(
                    resume=_truncate_for_ai(resume_text),
                    job=_truncate_for_ai(job_description),
                ),
                max_tokens=220,
                timeout=220.0,
            )
        )
        if ai_response.provider.value != "local":
            summary = ai_response.content.strip()
    except Exception:
        summary = ""

    if not summary:
        skills_text = ", ".join(matched) if matched else "relevant technical skills"
        summary = (
            f"Experienced professional with hands-on skills in {skills_text}, "
            "aligned with this role's requirements."
        )

    generated_text = f"{summary}\n\n---\n\n{resume_text}"

    file_path = TAILORED_DIR / f"{job_id}.docx"
    _write_docx(file_path, summary, resume_text)

    result = await db.execute(select(TailoredResume).where(TailoredResume.job_id == job_id))
    existing = result.scalar_one_or_none()

    if existing is not None:
        existing.summary = summary
        existing.generated_text = generated_text
        existing.file_path = str(file_path)
        existing.matched_skills = matched
        existing.missing_skills = missing
        existing.score = score
        tailored = existing
    else:
        tailored = TailoredResume(
            job_id=job_id,
            summary=summary,
            generated_text=generated_text,
            file_path=str(file_path),
            matched_skills=matched,
            missing_skills=missing,
            score=score,
        )
        db.add(tailored)

    await db.flush()
    await db.refresh(tailored)

    return tailored


async def get_tailored_resume(db: AsyncSession, job_id: int) -> TailoredResume | None:
    result = await db.execute(select(TailoredResume).where(TailoredResume.job_id == job_id))
    return result.scalar_one_or_none()
