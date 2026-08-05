import io
import uuid

import pytest
from docx import Document
from sqlalchemy import delete

from smarthunt.database.models.job import Job


async def _auth_headers(client) -> dict:
    uid = uuid.uuid4().hex[:8]
    payload = {
        "username": f"search_user_{uid}",
        "email": f"{uid}@example.com",
        "password": "Secret123",
    }
    await client.post("/api/v1/auth/register", json=payload)
    login = await client.post(
        "/api/v1/auth/login",
        json={"username": payload["username"], "password": payload["password"]},
    )
    return {"Authorization": f"Bearer {login.json()['access_token']}"}


@pytest.mark.asyncio
async def test_search_jobs_endpoint(client, db_session):
    """Job search hits the real database, not a hardcoded fixture list."""
    db_session.add(
        Job(
            title="Senior Backend Engineer Unique Marker",
            company="Acme",
            location="Remote",
            source="linkedin",
            url="https://example.com/jobs/senior-backend-engineer",
        )
    )
    await db_session.commit()

    # Filter by keyword rather than relying on default pagination order —
    # the jobs table is real and shared with live discovery runs (including
    # this same test DB), so unrelated jobs can easily outnumber the
    # default page size.
    response = await client.get(
        "/api/v1/search/jobs", params={"keyword": "Senior Backend Engineer Unique Marker"}
    )
    assert response.status_code == 200
    data = response.json()
    assert "jobs" in data
    assert "total" in data
    assert isinstance(data["jobs"], list)
    assert len(data["jobs"]) > 0
    assert any(job["title"] == "Senior Backend Engineer Unique Marker" for job in data["jobs"])


@pytest.mark.asyncio
async def test_search_score_reflects_real_resume_match(tmp_path, monkeypatch, client, db_session):
    """score_min/sort=score used to be accepted but silently ignored (no
    salary/score data existed anywhere). They now reflect a real match
    against the uploaded resume's text."""
    storage_dir = tmp_path / "resumes"
    storage_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr("smarthunt.resume.storage.storage.STORAGE_DIR", storage_dir)

    headers = await _auth_headers(client)

    doc_buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Experienced Python and Docker engineer.")
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    upload_response = await client.post(
        "/api/v1/resume/upload",
        files={
            "file": (
                "resume.docx",
                doc_buffer,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=headers,
    )
    assert upload_response.status_code == 200

    db_session.add_all(
        [
            Job(
                title="Score Test Python Engineer",
                company="Acme",
                location="Remote",
                source="test",
                url="https://example.com/jobs/score-test-python-engineer",
                requirements="python docker",
            ),
            Job(
                title="Score Test AIX Administrator",
                company="Acme",
                location="Remote",
                source="test",
                url="https://example.com/jobs/score-test-aix-admin",
                requirements="aix",
            ),
        ]
    )
    await db_session.commit()

    response = await client.get(
        "/api/v1/search/jobs",
        params={"sort": "score", "order": "desc", "score_min": 1, "limit": 100},
    )
    assert response.status_code == 200
    jobs = response.json()["jobs"]
    titles = [job["title"] for job in jobs]

    assert "Score Test Python Engineer" in titles
    assert "Score Test AIX Administrator" not in titles

    python_job = next(j for j in jobs if j["title"] == "Score Test Python Engineer")
    assert python_job["score"] == 100


@pytest.mark.asyncio
async def test_search_includes_score_without_explicit_sort(
    tmp_path, monkeypatch, client, db_session
):
    """Regression test: score used to only be computed when the caller
    explicitly sorted/filtered by it (sort=score, score_min, score_max),
    so the Jobs tab could only show a match-% column while that toggle
    was on. It's a cheap, rule-based computation (no AI call) — compute
    it on every search so the column can just always be there."""
    storage_dir = tmp_path / "resumes"
    storage_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr("smarthunt.resume.storage.storage.STORAGE_DIR", storage_dir)

    headers = await _auth_headers(client)

    doc_buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Experienced Python and Docker engineer.")
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    upload_response = await client.post(
        "/api/v1/resume/upload",
        files={
            "file": (
                "resume.docx",
                doc_buffer,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=headers,
    )
    assert upload_response.status_code == 200

    # This exact marker job may already exist from a prior run of this
    # same test (db_session.add() here bypasses the repository's own
    # dedup check) — clear it first so the assertion isn't order-dependent.
    await db_session.execute(delete(Job).where(Job.title == "Default Sort Score Marker Job"))
    await db_session.commit()

    db_session.add(
        Job(
            title="Default Sort Score Marker Job",
            company="Acme",
            location="Remote",
            source="test",
            url="https://example.com/jobs/default-sort-score-marker",
            requirements="python docker",
        )
    )
    await db_session.commit()

    response = await client.get(
        "/api/v1/search/jobs",
        params={"keyword": "Default Sort Score Marker Job"},
    )
    assert response.status_code == 200
    jobs = response.json()["jobs"]
    assert len(jobs) == 1
    assert jobs[0]["score"] == 100
